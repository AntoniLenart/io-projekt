import os
import ffmpeg
import whisper
from docx import Document


class Transcription:
    def __init__(self):
        self.model = whisper.load_model("base")

    def transcribe(self, video_path, output_dir):
        audio_path = os.path.join(output_dir, "audio.wav")
        ffmpeg.input(video_path).output(audio_path, ac=1, ar='16000').run(overwrite_output=True)
        result = self.model.transcribe(audio_path, language="polish")
        return result["text"]

    def generate_notes(self, transcription, output_dir):
        doc = Document()
        doc.add_heading("Streszczenie Spotkania", 0)
        doc.add_paragraph(transcription[:500] + "..." if len(transcription) > 500 else transcription)
        doc.add_heading("Transkrypcja", level=1)
        doc.add_paragraph(transcription)
        doc.save(os.path.join(output_dir, "notes.docx"))
